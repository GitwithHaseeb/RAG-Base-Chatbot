"""Generate sample PDF and DOCX files for GH Buddy."""

from __future__ import annotations

from pathlib import Path

from docx import Document
from reportlab.lib.pagesizes import LETTER
from reportlab.pdfgen import canvas


def create_pdf(path: Path) -> None:
    path.parent.mkdir(parents=True, exist_ok=True)
    c = canvas.Canvas(str(path), pagesize=LETTER)
    c.setFont("Helvetica-Bold", 16)
    c.drawString(72, 740, "Dummy Course Syllabus")
    c.setFont("Helvetica", 11)
    lines = [
        "Course: Introduction to Programming",
        "Week 1: Variables, Data Types, and Input/Output",
        "Week 2: Conditionals and Loops",
        "Week 3: Functions and Modules",
        "Week 4: Basic Data Structures",
        "Office hours every Tuesday at 3 PM.",
    ]
    y = 710
    for line in lines:
        c.drawString(72, y, line)
        y -= 22
    c.save()


def create_docx(path: Path) -> None:
    path.parent.mkdir(parents=True, exist_ok=True)
    doc = Document()
    doc.add_heading("Python Notes (Sample)", level=1)
    doc.add_paragraph("Python is an interpreted, high-level programming language.")
    doc.add_paragraph("Lists are mutable sequences used to store collections of items.")
    doc.add_paragraph("Use dictionaries for key-value mappings.")
    doc.add_paragraph("Always write readable code and meaningful function names.")
    doc.save(str(path))


if __name__ == "__main__":
    create_pdf(Path("data/syllabus.pdf"))
    create_docx(Path("data/python_notes.docx"))
    print("Sample documents created in data/ directory.")
