"""Generate a polished Word report from docs/project_report.md."""

from __future__ import annotations

from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Inches, Pt


ROOT = Path(__file__).parent
MD_PATH = ROOT / "docs" / "project_report.md"
DOCX_PATH = ROOT / "docs" / "project_report.docx"
LOGO_PATH = ROOT / "logo" / "gh_logo.png"


def add_title_page(doc: Document) -> None:
    if LOGO_PATH.exists():
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run()
        run.add_picture(str(LOGO_PATH), width=Inches(2.2))

    title = doc.add_paragraph("GH Buddy Chatbot")
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title.runs[0].font.size = Pt(24)
    title.runs[0].font.bold = True

    subtitle = doc.add_paragraph("Comprehensive Project Report")
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    subtitle.runs[0].font.size = Pt(14)

    authors = doc.add_paragraph("Authors: GH Team")
    authors.alignment = WD_ALIGN_PARAGRAPH.CENTER
    authors.runs[0].font.size = Pt(11)

    doc.add_page_break()


def add_markdown_content(doc: Document, text: str) -> None:
    in_code_block = False
    for raw_line in text.splitlines():
        line = raw_line.rstrip()
        stripped = line.strip()

        if stripped.startswith("```"):
            in_code_block = not in_code_block
            continue

        if in_code_block:
            p = doc.add_paragraph(line)
            p.style = "No Spacing"
            continue

        if not stripped:
            doc.add_paragraph("")
            continue

        if stripped.startswith("# "):
            doc.add_heading(stripped[2:], level=1)
            continue

        if stripped.startswith("## "):
            doc.add_heading(stripped[3:], level=2)
            continue

        if stripped.startswith("### "):
            doc.add_heading(stripped[4:], level=3)
            continue

        if stripped.startswith("- "):
            doc.add_paragraph(stripped[2:], style="List Bullet")
            continue

        if stripped.startswith("|") and stripped.endswith("|"):
            # Keep markdown table row as plain paragraph for compatibility.
            doc.add_paragraph(stripped)
            continue

        if stripped.startswith(("1. ", "2. ", "3. ", "4. ", "5. ", "6. ", "7. ", "8. ", "9. ")):
            doc.add_paragraph(stripped, style="List Number")
            continue

        doc.add_paragraph(stripped)


def main() -> None:
    if not MD_PATH.exists():
        raise FileNotFoundError(f"Missing markdown report: {MD_PATH}")

    markdown_text = MD_PATH.read_text(encoding="utf-8")
    doc = Document()
    add_title_page(doc)
    add_markdown_content(doc, markdown_text)
    DOCX_PATH.parent.mkdir(parents=True, exist_ok=True)
    doc.save(str(DOCX_PATH))
    print(f"Word report generated: {DOCX_PATH}")


if __name__ == "__main__":
    main()
